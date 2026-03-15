"""
Resume File Parser
Handles PDF, DOCX, JPG, PNG with OCR fallback
"""
import os
import logging
import re
from typing import Tuple

logger = logging.getLogger(__name__)


def parse_file(filepath: str, file_ext: str) -> Tuple[str, str]:
    """
    Parse a resume file and return (text, method_used).
    Returns (text, method) tuple.
    """
    ext = file_ext.lower().lstrip('.')

    if ext == 'pdf':
        return _parse_pdf(filepath)
    elif ext == 'docx':
        return _parse_docx(filepath)
    elif ext in ('jpg', 'jpeg', 'png'):
        return _parse_image(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(filepath: str) -> Tuple[str, str]:
    """Extract text from PDF using pdfplumber, fallback to OCR."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text = '\n'.join(pages_text)

        # Clean up extracted text
        text = _clean_pdf_text(text)

        if len(text.strip()) > 50:
            logger.info(f"PDF parsed with pdfplumber: {len(text)} chars")
            return text, 'pdfplumber'

    except ImportError:
        logger.warning("pdfplumber not available")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Fallback to OCR
    logger.info("PDF text extraction insufficient, trying OCR")
    return _pdf_ocr(filepath)


def _pdf_ocr(filepath: str) -> Tuple[str, str]:
    """Convert PDF pages to images and OCR them."""
    try:
        import fitz  # PyMuPDF
        import tempfile
        from PIL import Image
        import pytesseract

        doc = fitz.open(filepath)
        all_text = []

        for page_num in range(min(len(doc), 5)):  # Max 5 pages
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for quality

            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp_path = tmp.name
                pix.save(tmp_path)

            try:
                img = Image.open(tmp_path)
                text = pytesseract.image_to_string(img, config='--psm 6')
                all_text.append(text)
            finally:
                os.unlink(tmp_path)

        doc.close()
        return '\n'.join(all_text), 'ocr_pdf'

    except ImportError as e:
        logger.warning(f"OCR libraries not available: {e}")
        return "", 'failed'
    except Exception as e:
        logger.error(f"PDF OCR failed: {e}")
        return "", 'failed'


def _parse_docx(filepath: str) -> Tuple[str, str]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = '\n'.join(paragraphs)
        logger.info(f"DOCX parsed: {len(text)} chars")
        return text, 'python-docx'

    except ImportError:
        logger.warning("python-docx not available")
        return "", 'failed'
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return "", 'failed'


def _parse_image(filepath: str) -> Tuple[str, str]:
    """OCR an image file."""
    try:
        import pytesseract
        from PIL import Image, ImageEnhance, ImageFilter

        img = Image.open(filepath)

        # Preprocess for better OCR
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Enhance contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)

        text = pytesseract.image_to_string(img, config='--psm 6 --oem 3')
        logger.info(f"Image OCR completed: {len(text)} chars")
        return text, 'tesseract_ocr'

    except ImportError:
        logger.warning("pytesseract/PIL not available")
        return "", 'failed'
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return "", 'failed'


def _clean_pdf_text(text: str) -> str:
    """Clean PDF extraction artifacts."""
    if not text:
        return ""
    # Fix broken words from PDF column layout
    text = re.sub(r'([a-z])\s+([a-z])', lambda m: m.group(0)
                  if len(m.group(0).split()) > 3 else m.group(0), text)
    # Remove page numbers
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    # Remove excessive dots (table of contents artifacts)
    text = re.sub(r'\.{3,}', ' ', text)
    # Fix common PDF encoding issues
    text = text.replace('\x00', '').replace('\uf0b7', '•')
    return text
